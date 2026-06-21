"""Build and update a local Chroma knowledge base for personalized retrieval."""

import argparse
import hashlib
import shutil
from datetime import datetime, timezone
from pathlib import Path

from langchain_core.documents import Document
from langchain_core.runnables import RunnableConfig
from langchain_text_splitters import RecursiveCharacterTextSplitter

from open_deep_research.configuration import Configuration
from open_deep_research.rag import _build_vectorstore

SUPPORTED_SUFFIXES = {".md", ".txt", ".pdf", ".docx"}


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf_file(path: Path) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise ImportError(
            "PyMuPDF is required to ingest PDF files. Install with `pip install pymupdf`."
        ) from exc

    parts: list[str] = []
    with fitz.open(path) as pdf_doc:
        for page in pdf_doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to ingest DOCX files. Install with `pip install python-docx`."
        ) from exc

    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _load_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _read_text_file(path)
    if suffix == ".pdf":
        return _read_pdf_file(path)
    if suffix == ".docx":
        return _read_docx_file(path)
    return ""


def _collect_source_documents(source_dir: Path) -> list[Document]:
    docs: list[Document] = []
    for path in source_dir.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue

        text = _load_file(path).strip()
        if not text:
            continue

        stats = path.stat()
        source_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source_path": str(path.resolve()),
                    "doc_type": path.suffix.lower().lstrip("."),
                    "source_hash": source_hash,
                    "updated_at": datetime.fromtimestamp(
                        stats.st_mtime, tz=timezone.utc
                    ).isoformat(),
                },
            )
        )
    return docs


def ingest_local_knowledge_base(
    source: str,
    config: RunnableConfig | None = None,
    rebuild: bool = False,
) -> dict:
    """Ingest local files into Chroma and return indexing statistics."""
    configurable = Configuration.from_runnable_config(config)
    source_dir = Path(source).resolve()
    if not source_dir.exists() or not source_dir.is_dir():
        raise ValueError(f"Source directory does not exist: {source_dir}")

    persist_dir = Path(configurable.chroma_persist_directory).resolve()
    if rebuild and persist_dir.exists():
        shutil.rmtree(persist_dir)

    source_docs = _collect_source_documents(source_dir)
    if not source_docs:
        return {
            "status": "no_documents",
            "source_files": 0,
            "chunks": 0,
            "persist_directory": str(persist_dir),
        }

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    chunked_docs = splitter.split_documents(source_docs)
    if not chunked_docs:
        return {
            "status": "no_chunks",
            "source_files": len(source_docs),
            "chunks": 0,
            "persist_directory": str(persist_dir),
        }

    vectorstore = _build_vectorstore(configurable, config)

    # Keep the index idempotent: replace existing chunks for each source file.
    source_paths = {doc.metadata.get("source_path") for doc in source_docs}
    for source_path in source_paths:
        if source_path:
            vectorstore.delete(where={"source_path": source_path})

    ids: list[str] = []
    for idx, doc in enumerate(chunked_docs):
        source_path = str(doc.metadata.get("source_path", "unknown"))
        source_hash = str(doc.metadata.get("source_hash", ""))
        doc_id = hashlib.sha256(
            f"{source_path}:{source_hash}:{idx}".encode("utf-8")
        ).hexdigest()
        ids.append(doc_id)

    vectorstore.add_documents(chunked_docs, ids=ids)

    return {
        "status": "ok",
        "source_files": len(source_docs),
        "chunks": len(chunked_docs),
        "persist_directory": str(persist_dir),
    }


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Build local knowledge base index for RAG.")
    parser.add_argument(
        "--source",
        required=True,
        help="Path to local knowledge folder containing md/txt/pdf/docx files.",
    )
    parser.add_argument(
        "--rebuild",
        action="store_true",
        help="If set, remove existing Chroma persistence directory before indexing.",
    )
    return parser


def main() -> None:
    parser = _build_arg_parser()
    args = parser.parse_args()
    result = ingest_local_knowledge_base(source=args.source, rebuild=args.rebuild)
    print(result)
    if result.get("status") in {"no_documents", "no_chunks"}:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
